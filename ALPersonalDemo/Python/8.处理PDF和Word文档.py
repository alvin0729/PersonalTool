

13.1.1 从PDF提取文本

>>> import PyPDF2
>>> pdfFileObj = open('meetingminutes.pdf', 'rb')
>>> pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
>>> pdfReader.numPages
19
>>> pageObj = pdfReader.getPage(0)
>>> pageObj.extractText()
'OOFFFFIICCIIAALL BBOOAARRDD MMIINNUUTTEESS Meeting of March 7, 2015


13.1.2 解密PDF
>>> import PyPDF2
>>> pdfReader = PyPDF2.PdfFileReader(open('encrypted.pdf', 'rb'))
>>> pdfReader.isEncrypted
True
>>> pdfReader.getPage(0)
>>> pdfReader.decrypt('rosebud')
1
>>> pageObj = pdfReader.getPage(0)

13.1.3 创建PDF
PyPDF2 写入 PDF 的能力，仅限于从其他 PDF 中拷贝页面、旋转
页面、重叠页面和加密文件。


13.1.4 拷贝页面（合并页面）
>>> import PyPDF2
>>> pdf1File = open('meetingminutes.pdf', 'rb')
>>> pdf2File = open('meetingminutes2.pdf', 'rb')
>>> pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
>>> pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
>>> pdfWriter = PyPDF2.PdfFileWriter()
>>> for pageNum in range(pdf1Reader.numPages):
        pageObj = pdf1Reader.getPage(pageNum)
        pdfWriter.addPage(pageObj)
>>> for pageNum in range(pdf2Reader.numPages):
        pageObj = pdf2Reader.getPage(pageNum)
        pdfWriter.addPage(pageObj)
>>> pdfOutputFile = open('combinedminutes.pdf', 'wb')
>>> pdfWriter.write(pdfOutputFile)
>>> pdfOutputFile.close()
>>> pdf1File.close()
>>> pdf2File.close()


13.1.5 旋转页面
>>> import PyPDF2
>>> minutesFile = open('meetingminutes.pdf', 'rb')
>>> pdfReader = PyPDF2.PdfFileReader(minutesFile)
>>> page = pdfReader.getPage(0)
>>> page.rotateClockwise(90)
{'/Contents': [IndirectObject(961, 0), IndirectObject(962, 0), IndirectObject(963, 0), IndirectObject(964, 0), IndirectObject(965, 0), IndirectObject(966, 0), IndirectObject(967, 0), IndirectObject(968, 0)], '/CropBox': [0, 0, 612, 792], '/MediaBox': [0, 0, 612, 792], '/Parent': IndirectObject(953, 0), '/Resources': {'/ColorSpace': {'/CS0': IndirectObject(975, 0), '/CS1': IndirectObject(976, 0), '/CS2': IndirectObject(976, 0)}, '/ExtGState': {'/GS0': IndirectObject(977, 0)}, '/Font': {'/TT0': IndirectObject(979, 0), '/TT1': IndirectObject(981, 0), '/TT2': IndirectObject(983, 0), '/TT3': IndirectObject(985, 0), '/TT4': IndirectObject(987, 0), '/TT5': IndirectObject(989, 0)}, '/XObject': {'/Im0': IndirectObject(973, 0)}}, '/Rotate': 90, '/StructParents': 0, '/Type': '/Page'}
>>> pdfWriter = PyPDF2.PdfFileWriter()
>>> pdfWriter.addPage(page)
>>> resultPdfFile = open('rotatedPage.pdf', 'wb')
>>> pdfWriter.write(resultPdfFile)
>>> resultPdfFile.close()
>>> minutesFile.close()


13.1.6 叠加页面
>>> import PyPDF2
>>> minutesFile = open('meetingminutes.pdf', 'rb')
>>> pdfReader = PyPDF2.PdfFileReader(minutesFile)
>>> minutesFirstPage = pdfReader.getPage(0)
>>> pdfWatermarkReader = PyPDF2.PdfFileReader(open('watermark.pdf', 'rb'))
>>> minutesFirstPage.mergePage(pdfWatermarkReader.getPage(0))
>>> pdfWriter = PyPDF2.PdfFileWriter()
>>> pdfWriter.addPage(minutesFirstPage)
>>> for pageNum in range(1, pdfReader.numPages):
...     pageObj = pdfReader.getPage(pageNum)
...     pdfWriter.addPage(pageObj)
... 
>>> resultPdfFile = open('watermarkedCover.pdf', 'wb')
>>> pdfWriter.write(resultPdfFile)
>>> minutesFile.close()
>>> resultPdfFile.close()


13.1.7 加密PDF
>>> import PyPDF2
>>> pdfFile = open('meetingminutes.pdf', 'rb')
>>> pdfReader = PyPDF2.PdfFileReader(pdfFile)
>>> pdfWriter = PyPDF2.PdfFileWriter()
>>> for pageNum in range(pdfReader.numPages):
...     pdfWriter.addPage(pdfReader.getPage(pageNum))
... 
>>> pdfWriter.encrypt('swordfish')
>>> resultPdf = open('encryptedminutes.pdf', 'wb')
>>> pdfWriter.write(resultPdf)
>>> resultPdf.close()


13.2 项目：从多个PDF中合并选择的页面
调用 os.listdir()，找到当前工作目录中的所有文件，去除掉非 PDF 文件。
调用 Python 的 sort()列表方法，对文件名按字母排序。
为输出的 PDF 文件创建 PdfFileWriter 对象。
循环遍历每个 PDF 文件， 为它创建 PdfFileReader 对象。
针对每个 PDF 文件，循环遍历每一页，第一页除外。
将页面添加到输出的 PDF。
将输出的 PDF 写入一个文件，名为 allminutes.pdf。

#! python3
# combinePdfs.py - Combines all the PDFs in the current working directory into 
# a single PDF.

import PyPDF2, os

# Get all the PDF filenames.
pdfFiles = []
for filename in os.listdir('.'):
    if filename.endswith('.pdf'):
        pdfFiles.append(filename)
pdfFiles.sort()

pdfWriter = PyPDF2.PdfFileWriter()

# Loop through all the PDF files.
for filename in pdfFiles:
    pdfFileObj = open(filename, 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

    # Loop through all the pages (except the first) and add them.
    for pageNum in range(1, pdfReader.numPages):
        pageObj = pdfReader.getPage(pageNum)
        pdfWriter.addPage(pageObj)

# Save the resulting PDF to a file.
pdfOutput = open('allminutes.pdf', 'wb')
pdfWriter.write(pdfOutput)
pdfOutput.close()


13.3 Word文档
pip install python-docx

在最高一层， Document 对象表示整个文档
Document 对象包含一个 Paragraph 对象的列表，表示文档中的段落
每个 Paragraph 对象都包含一个 Run 对象的列表(一个 Run 对象是相同样式文本的延续。当文本样式发生改变时，就需要一个新的 Run 对象。)

13.3.1 读取Word文档
>>> import docx
>>> doc = docx.Document('demo.docx')
>>> len(doc.paragraphs)
7
>>> doc.paragraphs[0].text
'Document Title'
>>> doc.paragraphs[1].text
'A plain paragraph with some bold and some italic'
>>> len(doc.paragraphs[1].runs)
5
>>> doc.paragraphs[1].runs[0].text
'A plain paragraph with'
>>> doc.paragraphs[1].runs[1].text
' some '
>>> doc.paragraphs[1].runs[2].text
'bold'
>>> doc.paragraphs[1].runs[3].text
' and some '
>>> doc.paragraphs[1].runs[4].text
'italic'

13.3.2 从.docx文件中取得完整的文本
#! python3

import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n\n'.join(fullText)
print(getText('demo.docx'))


13.3.3 设置Paragraph 和 Run对象的样式
>>> doc = docx.Document('demo.docx')
>>> doc.paragraphs[0].text
'Document Title'
>>> doc.paragraphs[0].style
_ParagraphStyle('Title') id: 4354126904
>>> doc.paragraphs[0].style = 'Normal'
>>> doc.paragraphs[1].text
'A plain paragraph with some bold and some italic'
>>> (doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.
paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
... ('A plain paragraph with', ' some ', 'bold', ' and some ')
>>> doc.paragraphs[1].runs[0].style = 'QuoteChar'
>>> doc.paragraphs[1].runs[1].underline = True
>>> doc.paragraphs[1].runs[3].underline = True
>>> doc.save('restyled.docx')


13.3.6 写入Word文档
>>> import docx
>>> doc = docx.Document()
>>> doc.add_paragraph('Hello world!')
<docx.text.paragraph.Paragraph object at 0x10387eb00>
>>> doc.save('helloworld.docx')


>>> doc = docx.Document()
>>> doc.add_paragraph('Hello world!')
<docx.text.paragraph.Paragraph object at 0x103898550>
>>> paraObj1 = doc.add_paragraph('This is a second paragraph.')
>>> paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
>>> paraObj1.add_run(' This text is being added to the second paragraph.')
<docx.text.run.Run object at 0x1038984a8>
>>> doc.save('multipleParagraphs.docx')

add_paragraph()和 add_run()都接受可选的第二个参数，它是表示 Paragraph 或 Run 对象样式的字符串。
doc.add_paragraph('Hello world!', 'Title')

13.3.7 添加标题
>>> doc = docx.Document()
>>> doc.add_heading('Header 0', 0)
<docx.text.paragraph.Paragraph object at 0x103855f98>
>>> doc.add_heading('Header 1', 1)
<docx.text.paragraph.Paragraph object at 0x10227c438>
>>> doc.add_heading('Header 2', 2)
<docx.text.paragraph.Paragraph object at 0x103855d30>
>>> doc.add_heading('Header 3', 3)
<docx.text.paragraph.Paragraph object at 0x10227c438>
>>> doc.add_heading('Header 4', 4)
<docx.text.paragraph.Paragraph object at 0x103855f98>
>>> doc.save('headings.docx')


13.3.8 添加换行符和换页符
>>> doc = docx.Document()
>>> doc.add_paragraph('This is on the first page!')
<docx.text.Paragraph object at 0x0000000003785518>
n >>> doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE)
>>> doc.add_paragraph('This is on the second page!')
<docx.text.Paragraph object at 0x00000000037855F8>
>>> doc.save('twoPage.docx')


13.3.9 添加图像
>>> doc.add_picture('zophie.png', width=docx.shared.Inches(1),
height=docx.shared.Cm(4))
<docx.shape.InlineShape object at 0x00000000036C7D30>

















